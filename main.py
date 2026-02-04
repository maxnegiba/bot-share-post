import os
import sys
import time
import math
import logging
import random
import re
import io
import sqlite3
import secrets
import pickle
from datetime import datetime, timedelta
from typing import Optional, List
from logging.handlers import RotatingFileHandler

# Third-party Utils
from dotenv import load_dotenv
import psutil
from tenacity import retry, stop_after_attempt, wait_fixed

# Parsing Imports
import docx
import pdfplumber

# Google Imports
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account

# Undetected Selenium Imports
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException, 
    NoSuchElementException, 
    StaleElementReferenceException,
    ElementClickInterceptedException,
    WebDriverException
)

# --- 1. CONFIGURARE & LOGGING ---
load_dotenv()

log_formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
file_handler = RotatingFileHandler(
    "bot_stealth_v12.log", 
    maxBytes=10*1024*1024, 
    backupCount=5, 
    encoding='utf-8'
)
file_handler.setFormatter(log_formatter)
stream_handler = logging.StreamHandler()
stream_handler.setFormatter(log_formatter)
logging.basicConfig(
    level=logging.INFO, 
    handlers=[file_handler, stream_handler]
)

class Config:
    FB_EMAIL = os.getenv('FB_EMAIL')
    FB_PASSWORD = os.getenv('FB_PASSWORD')
    FB_PAGE_NAME = os.getenv('FB_PAGE_NAME')
    PROXY_URL = os.getenv('PROXY_URL', None)
    PROXY_TIMEZONE = os.getenv('PROXY_TIMEZONE', 'Europe/Bucharest')
    GOOGLE_JSON = os.getenv('GOOGLE_CREDENTIALS_FILE', 'service_account.json')
    DRIVE_DOCX_NAME = os.getenv('DRIVE_FILENAME_DOCX')
    DRIVE_PDF_NAME = os.getenv('DRIVE_FILENAME_PDF')
    TEMP_DIR = os.getenv('TEMP_DOWNLOAD_DIR', './temp_downloads')
    LOCAL_PATH_DOCX = os.path.join(TEMP_DIR, 'posts_schedule.docx')
    LOCAL_PATH_PDF = os.path.join(TEMP_DIR, 'groups_list.pdf')
    DB_PATH = "history_v12.db"
    PROFILE_PATH = os.path.abspath(os.getenv('CHROME_PROFILE_PATH', './chrome_data'))
    
    # --- MODIFICARE: Eliminated Start/End Hour ---
    # Botul va rula continuu in limitele zilnice
    
    DAILY_GROUP_LIMIT = int(os.getenv('DAILY_GROUP_LIMIT', 40))
    MAX_GROUPS_POOL = int(os.getenv('MAX_GROUPS_POOL', 150))
    DELAY_MIN_SEC = int(os.getenv('DELAY_MIN_SEC', 180))
    DELAY_MAX_SEC = int(os.getenv('DELAY_MAX_SEC', 400))
    DO_WARMUP = os.getenv('DO_WARMUP', 'True').lower() == 'true'
    MAX_RETRIES = int(os.getenv('MAX_RETRIES', 3))
    SESSION_REFRESH_HOURS = int(os.getenv('SESSION_REFRESH_HOURS', 4))
    COOKIE_PATH = "fb_session.pkl"

    @classmethod
    def validate(cls):
        required = {
            'FB_EMAIL': cls.FB_EMAIL,
            'FB_PASSWORD': cls.FB_PASSWORD,
            'DRIVE_DOCX_NAME': cls.DRIVE_DOCX_NAME,
            'DRIVE_PDF_NAME': cls.DRIVE_PDF_NAME
        }
        missing = [k for k, v in required.items() if not v]
        if missing:
            raise ValueError(f"❌ Lipsesc variabile: {', '.join(missing)}")
        if not os.path.exists(cls.GOOGLE_JSON):
            raise FileNotFoundError(f"❌ Credentials inexistent: {cls.GOOGLE_JSON}")
        if cls.PROXY_URL:
            logging.info(f"🛡️ Proxy activat: {cls.PROXY_URL.split('@')[-1]}")
        else:
            logging.warning("⚠️ Fără Proxy (risc de IP ban)")
        logging.info("✓ Configurare validată")

# --- 2. DATABASE MANAGER ---
class DatabaseManager:
    def __init__(self, db_name=Config.DB_PATH):
        self.conn = sqlite3.connect(db_name, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self._init_db()

    def _init_db(self):
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS posted_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                post_date TEXT NOT NULL,
                link_url TEXT NOT NULL,
                link_base_url TEXT NOT NULL,
                group_url TEXT NOT NULL,
                status TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                retry_count INTEGER DEFAULT 0
            )
        ''')
        self.cursor.execute('CREATE INDEX IF NOT EXISTS idx_date_group ON posted_history(post_date, group_url)')
        self.cursor.execute('CREATE INDEX IF NOT EXISTS idx_base_url ON posted_history(link_base_url)')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS app_state (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        ''')
        self.conn.commit()
        logging.info("✓ Database inițializat")

    def get_group_index(self) -> int:
        try:
            self.cursor.execute("SELECT value FROM app_state WHERE key = 'last_group_index'")
            row = self.cursor.fetchone()
            return int(row[0]) if row else 0
        except:
            return 0

    def update_group_index(self, new_index: int):
        self.cursor.execute("""
            INSERT INTO app_state (key, value) VALUES ('last_group_index', ?)
            ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """, (str(new_index),))
        self.conn.commit()

    def already_posted_today(self, group_url: str, link_base_url: str) -> bool:
        today = datetime.now().strftime("%Y-%m-%d")
        self.cursor.execute('''
            SELECT id FROM posted_history 
            WHERE post_date = ? 
            AND group_url = ? 
            AND link_base_url = ?
            AND status IN ('SUCCESS', 'SUCCESS_FALLBACK')
        ''', (today, group_url, link_base_url))
        return self.cursor.fetchone() is not None

    def log_post(self, link_url: str, link_base_url: str, group_url: str, status: str, retry_count: int = 0):
        today = datetime.now().strftime("%Y-%m-%d")
        self.cursor.execute('''
            INSERT INTO posted_history 
            (post_date, link_url, link_base_url, group_url, status, retry_count) 
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (today, link_url, link_base_url, group_url, status, retry_count))
        self.conn.commit()

    def get_today_stats(self) -> dict:
        today = datetime.now().strftime("%Y-%m-%d")
        self.cursor.execute('''
            SELECT status, COUNT(*) 
            FROM posted_history 
            WHERE post_date = ? 
            GROUP BY status
        ''', (today,))
        return dict(self.cursor.fetchall())

    def cleanup_old_records(self, days: int = 30):
        cutoff = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
        self.cursor.execute('DELETE FROM posted_history WHERE post_date < ?', (cutoff,))
        deleted = self.cursor.rowcount
        self.conn.commit()
        if deleted > 0:
            logging.info(f"🗑️ Șterse {deleted} înregistrări > {days} zile")

    def __del__(self):
        if hasattr(self, 'conn'):
            self.conn.close()

# --- 3. GOOGLE DRIVE DOWNLOADER ---
class DriveDownloader:
    def __init__(self):
        self.scopes = ['https://www.googleapis.com/auth/drive.readonly']
        self.service = None
        self._authenticate()

    def _authenticate(self):
        try:
            if not os.path.exists(Config.GOOGLE_JSON):
                raise FileNotFoundError(f"❌ Fisierul {Config.GOOGLE_JSON} nu exista!")
            creds = service_account.Credentials.from_service_account_file(
                Config.GOOGLE_JSON, scopes=self.scopes
            )
            self.service = build('drive', 'v3', credentials=creds)
            logging.info("✓ Google Drive Service conectat")
        except Exception as e:
            logging.critical(f"❌ Eroare autentificare Drive: {e}")
            raise

    def _find_file_id(self, file_name: str) -> Optional[str]:
        try:
            query = f"name = '{file_name}' and trashed = false"
            results = self.service.files().list(q=query, pageSize=1, fields="files(id, name)").execute()
            items = results.get('files', [])
            if not items:
                logging.warning(f"⚠️ Fișierul '{file_name}' nu a fost găsit pe Drive.")
                return None
            return items[0]['id']
        except Exception as e:
            logging.error(f"❌ Eroare căutare fișier {file_name}: {e}")
            return None

    def _download_content(self, file_id: str, local_path: str):
        try:
            request = self.service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            with open(local_path, 'wb') as f:
                f.write(fh.getbuffer())
            logging.info(f"✓ Descărcat cu succes: {local_path}")
        except Exception as e:
            logging.error(f"❌ Eroare la descărcarea ID {file_id}: {e}")

    def sync_files(self):
        logging.info("📥 Încep sincronizarea fișierelor din Drive...")
        os.makedirs(Config.TEMP_DIR, exist_ok=True)

        docx_id = self._find_file_id(Config.DRIVE_DOCX_NAME)
        if docx_id:
            self._download_content(docx_id, Config.LOCAL_PATH_DOCX)
        else:
            logging.error(f"❌ Nu pot descărca {Config.DRIVE_DOCX_NAME} (ID negăsit)")

        pdf_id = self._find_file_id(Config.DRIVE_PDF_NAME)
        if pdf_id:
            self._download_content(pdf_id, Config.LOCAL_PATH_PDF)
        else:
            logging.error(f"❌ Nu pot descărca {Config.DRIVE_PDF_NAME} (ID negăsit)")

# --- 4. LOCAL PARSER ---
class LocalParser:
    @staticmethod
    def get_groups() -> List[str]:
        if not os.path.exists(Config.LOCAL_PATH_PDF):
            logging.warning(f"⚠️ PDF inexistent: {Config.LOCAL_PATH_PDF}")
            return []
            
        group_names = []
        try:
            with pdfplumber.open(Config.LOCAL_PATH_PDF) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if not text:
                        continue
                        
                    lines = text.split('\n')
                    for line in lines:
                        match = re.search(r"^\d+\.\s+(.+?)(?:\s+\d+[\.,]\d+)?$", line.strip())
                        if match:
                            raw_name = match.group(1)
                            clean_name = re.sub(r'\s*\d+[\.,]\d+\s*$', '', raw_name)
                            clean_name = clean_name.replace("...", "").strip()
                            
                            if len(clean_name) > 3 and clean_name not in group_names:
                                group_names.append(clean_name)

            logging.info(f"✓ Găsite {len(group_names)} nume de grupuri în PDF")
            if group_names:
                logging.info(f"   Exemple: {group_names[:3]}")
                
            return group_names

        except Exception as e:
            logging.error(f"❌ Eroare parsare PDF: {e}")
            return []

    @staticmethod
    def get_todays_post() -> Optional[str]:
        if not os.path.exists(Config.LOCAL_PATH_DOCX):
            return None
        try:
            doc = docx.Document(Config.LOCAL_PATH_DOCX)
            all_text = []
            for p in doc.paragraphs:
                all_text.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            all_text.append(p.text)
            
            urls = []
            pattern = r'https?://(?:www\.|m\.|web\.)?facebook\.com/[^\s]+'
            full_text = "\n".join(all_text)
            matches = re.findall(pattern, full_text)
            
            urls = [m.rstrip('.,;') for m in matches if "groups/" not in m]
            
            day = datetime.now().day
            if urls:
                index = (day - 1) % len(urls)
                return urls[index]
            return None
        except Exception as e:
            logging.error(f"❌ Eroare parsare DOCX: {e}")
            return None

# --- 5. UTILS ---
def kill_chrome():
    killed = 0
    for proc in psutil.process_iter(['name']):
        try:
            if 'chrome' in proc.info['name'].lower():
                proc.kill()
                killed += 1
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            pass
    if killed > 0:
        logging.info(f"🔪 Închise {killed} procese Chrome")
        time.sleep(2)

def generate_unique_link(base_link: str) -> str:
    unique_id = secrets.token_hex(4)
    timestamp = int(time.time())
    separator = "&" if "?" in base_link else "?"
    return f"{base_link}{separator}ref={unique_id}&ts={timestamp}"

def extract_base_link(full_link: str) -> str:
    return full_link.split('?')[0].split('#')[0]

# --- 6. FACEBOOK BOT ---
class FacebookBot:
    def __init__(self):
        self.driver = None
        self.wait = None
        self.actions = None
        self.session_start = datetime.now()
        self.selected_config = None
        self.viewport_width = None
        self.viewport_height = None
        self._init_driver()

    def _init_driver(self):
        kill_chrome()
        options = uc.ChromeOptions()
        options.add_argument(f"--user-data-dir={Config.PROFILE_PATH}")
        if Config.PROXY_URL:
            options.add_argument(f'--proxy-server={Config.PROXY_URL}')
        
        options.add_argument("--disable-notifications")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-infobars")
        options.add_argument("--disable-popup-blocking")

        ua_configs = [
            {"os": "Windows", "ua": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36", "platform": "Win32"},
            {"os": "Windows", "ua": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36", "platform": "Win32"},
            {"os": "MacOS", "ua": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36", "platform": "MacIntel"}
        ]
        self.selected_config = random.choice(ua_configs)
        options.add_argument(f"user-agent={self.selected_config['ua']}")
        logging.info(f"🎭 Profile Spoof: {self.selected_config['os']} | {self.selected_config['platform']}")

        try:
            self.driver = uc.Chrome(options=options, use_subprocess=True, version_main=None)
            self.wait = WebDriverWait(self.driver, 25)
            self.actions = ActionChains(self.driver)

            self.viewport_width = random.randint(1280, 1600)
            self.viewport_height = random.randint(800, 1000)
            self.driver.set_window_size(self.viewport_width, self.viewport_height)

            self._inject_stealth_scripts()
            if Config.PROXY_URL:
                self._patch_webrtc_leak()
            
            logging.info(f"✓ Driver init (Window: {self.viewport_width}x{self.viewport_height})")
        except Exception as e:
            logging.critical(f"❌ EROARE CRITICĂ DRIVER: {e}")
            raise

    # --- MANAGEMENT COOKIES & SESIUNE ---

    def save_cookies(self):
        try:
            cookies = self.driver.get_cookies()
            if not cookies:
                logging.warning("⚠️ Nu sunt cookie-uri de salvat.")
                return
            with open(Config.COOKIE_PATH, "wb") as f:
                pickle.dump(cookies, f)
            logging.info(f"🍪 {len(cookies)} Cookie-uri salvate în {Config.COOKIE_PATH}")
        except Exception as e:
            logging.warning(f"⚠️ Eroare salvare cookie-uri: {e}")

    def load_cookies(self):
        if not os.path.exists(Config.COOKIE_PATH):
            return False
        logging.info("🍪 Încerc restaurarea sesiunii din cookie-uri...")
        try:
            if "facebook.com" not in self.driver.current_url:
                self.driver.get("https://www.facebook.com/")
                time.sleep(3)
            with open(Config.COOKIE_PATH, "rb") as f:
                cookies = pickle.load(f)
            
            added = 0
            for cookie in cookies:
                if 'expiry' in cookie:
                    cookie['expiry'] = int(cookie['expiry'])
                try:
                    self.driver.add_cookie(cookie)
                    added += 1
                except: pass
            
            if added > 0:
                self.driver.refresh()
                time.sleep(5)
                return True
            return False
        except Exception as e:
            logging.error(f"❌ Eroare încărcare cookie-uri: {e}")
            return False

    def close(self):
        if self.driver:
            try:
                logging.info("🛑 Închidere driver...")
                self.driver.quit()
            except Exception:
                pass

    def refresh_session(self):
        try:
            self.driver.refresh()
            time.sleep(5)
            self._handle_cookie_consent()
        except:
            pass

    def is_session_expired(self) -> bool:
        try:
            url = self.driver.current_url.lower()
            if "login" in url or "checkpoint" in url:
                logging.warning("⚠️ URL suspect detectat (login/checkpoint).")
                return True
            
            login_form = self.driver.find_elements(By.ID, "login_form") or \
                         self.driver.find_elements(By.NAME, "login") or \
                         self.driver.find_elements(By.XPATH, "//input[@name='email']")
            
            if login_form and len(login_form) > 0:
                if self.driver.find_elements(By.XPATH, "//div[@role='banner']"):
                    return False
                return True
                
            return False
        except Exception:
            return True

    # --- DETECȚIE & LOGIN ---

    def ensure_logged_in(self) -> bool:
        try:
            logging.info("🔍 Verific starea sesiunii (Deep Check)...")
            
            if "facebook.com" not in self.driver.current_url:
                self.driver.get("https://www.facebook.com/")
                time.sleep(random.uniform(4, 6))

            self._handle_cookie_consent()

            logged_in_selectors = [
                "//div[@role='banner']",
                "//div[@role='feed']",
                "//div[@aria-label='Meniu' or @aria-label='Menu']", 
                "//div[@aria-label='Messenger']",
                "//div[@role='navigation']",
                "//svg[@aria-label='Home' or @aria-label='Acasă']",
                "//a[@aria-label='Facebook']"
            ]

            for xpath in logged_in_selectors:
                elements = self.driver.find_elements(By.XPATH, xpath)
                if elements and elements[0].is_displayed():
                    logging.info(f"✓ Sesiune activă confirmată (Selector: {xpath})")
                    return True

            login_inputs = self.driver.find_elements(By.ID, "email") or \
                           self.driver.find_elements(By.NAME, "email")
            
            if login_inputs:
                logging.info("ℹ️ Formular de login detectat. Inițiez autentificarea.")
                return self._perform_login()

            logging.warning("⚠️ Stare incertă (nici Feed, nici Login). Încerc refresh...")
            self.driver.refresh()
            time.sleep(8)
            
            if self.driver.find_elements(By.XPATH, "//div[@role='banner']"):
                logging.info("✓ Sesiune detectată după refresh.")
                return True

            return self._perform_login()

        except Exception as e:
            logging.error(f"❌ Eroare la verificarea sesiunii: {e}")
            return False

    def _perform_login(self) -> bool:
        logging.info("🔐 Încep procedura de login explicit...")
        try:
            self._handle_cookie_consent()

            try:
                email_field = self.wait.until(EC.presence_of_element_located((By.ID, "email")))
                email_field.clear()
                self.human_typing(email_field, Config.FB_EMAIL)
                time.sleep(1)
            except TimeoutException:
                if self.driver.find_elements(By.XPATH, "//div[@role='banner']"):
                    logging.info("✓ Login detectat (fără a introduce datele).")
                    return True
                logging.warning("⚠️ Nu găsesc câmpul email, dar nici nu par logat.")

            try:
                pass_field = self.driver.find_element(By.ID, "pass")
                pass_field.clear()
                self.human_typing(pass_field, Config.FB_PASSWORD)
                time.sleep(0.5)
                pass_field.send_keys(Keys.ENTER)
                logging.info("⏳ Credențiale trimise. Aștept încărcarea/2FA...")
            except Exception as e:
                logging.warning(f"⚠️ Problemă la introducerea parolei: {e}")

            max_wait_seconds = 300 
            start_wait = time.time()
            
            while time.time() - start_wait < max_wait_seconds:
                try:
                    if self.driver.find_elements(By.XPATH, "//div[@role='banner']") or \
                       self.driver.find_elements(By.XPATH, "//div[@role='feed']"):
                        logging.info("✓ Login reușit și confirmat!")
                        self.save_cookies()
                        return True
                    
                    save_btns = self.driver.find_elements(By.XPATH, "//div[contains(@aria-label, 'Save') or contains(text(), 'Save') or contains(text(), 'Salvează')]//span")
                    if save_btns:
                        for btn in save_btns:
                            if btn.is_displayed():
                                try: btn.click(); logging.info("ℹ️ Click 'Save Browser'"); break
                                except: pass
                    
                    cont_btns = self.driver.find_elements(By.XPATH, "//div[@role='button']//span[contains(text(), 'Continue')]")
                    if cont_btns:
                         try: cont_btns[0].click(); logging.info("ℹ️ Click 'Continue'"); break
                         except: pass

                except:
                    pass
                time.sleep(5)

            logging.error("❌ Timeout la login. Nu am detectat intrarea în cont.")
            self.driver.save_screenshot("login_timeout_debug.png")
            return False

        except Exception as e:
            logging.error(f"❌ Eroare fatală în procesul de login: {e}")
            return False

    # --- SCHIMBARE PROFIL ---

    def switch_profile(self):
        if not Config.FB_PAGE_NAME:
            logging.info("ℹ️ Nu este setat niciun nume de pagină. Rămân pe profil.")
            return

        logging.info(f"🔄 Verific profilul curent vs. Pagina dorită: '{Config.FB_PAGE_NAME}'")
        
        try:
            left_sidebar = self.driver.find_elements(By.XPATH, f"//div[@role='navigation']//span[contains(text(), '{Config.FB_PAGE_NAME}')]")
            avatar_lbl = self.driver.find_elements(By.XPATH, f"//div[@role='banner']//div[contains(@aria-label, '{Config.FB_PAGE_NAME}')]")
            
            if left_sidebar or avatar_lbl:
                logging.info(f"✅ SUNT DEJA PE PAGINA: {Config.FB_PAGE_NAME}. Nu fac switch.")
                return
        except Exception as e:
            logging.warning(f"⚠️ Eroare verificare profil curent: {e}")

        logging.info(f"🔄 Inițiez schimbarea pe: {Config.FB_PAGE_NAME}")
        try:
            banner_avatars = self.driver.find_elements(By.XPATH, "//div[@role='banner']//div[@role='button']//img")
            if banner_avatars:
                avatar_btn = banner_avatars[0].find_element(By.XPATH, "./../..")
                self.driver.execute_script("arguments[0].click();", avatar_btn)
                time.sleep(3)
                
                page_btn = self.driver.find_elements(By.XPATH, f"//span[contains(text(), '{Config.FB_PAGE_NAME}')]")
                if page_btn:
                    logging.info("✓ Găsit pagina în meniul principal. Click.")
                    page_btn[0].click()
                    time.sleep(8)
                    return
                
                see_all = self.driver.find_elements(By.XPATH, "//span[contains(text(), 'See all profiles') or contains(text(), 'Vezi toate')]")
                if see_all:
                    see_all[0].click()
                    time.sleep(3)
                    
                    page_btn_2 = self.driver.find_elements(By.XPATH, f"//span[contains(text(), '{Config.FB_PAGE_NAME}')]")
                    if page_btn_2:
                        logging.info("✓ Găsit pagina în 'See all'. Click.")
                        page_btn_2[0].click()
                        time.sleep(8)
                        return
        except Exception as e:
            logging.warning(f"⚠️ Eroare la switch prin meniu: {e}")

        logging.warning("⚠️ Metoda Meniu a eșuat. Încerc URL direct /pages...")
        self.driver.get("https://www.facebook.com/pages/?category=your_pages")
        time.sleep(6)
        
        try:
            switch_btn = self.driver.find_element(By.XPATH, f"//span[contains(text(), '{Config.FB_PAGE_NAME}')]/ancestor::div[contains(@role, 'article') or contains(@class, 'x')]//div[@role='button']")
            self.driver.execute_script("arguments[0].click();", switch_btn)
            time.sleep(8)
            logging.info("✓ Switch realizat din lista de pagini.")
        except:
            logging.error(f"❌ EȘEC TOTAL la schimbarea pe pagina: {Config.FB_PAGE_NAME}")

    # --- POSTARE ÎN GRUP ---
    def post_to_group(self, group_name: str, post_link: str) -> str:
        logging.info(f"🚀 Procesez grupul: '{group_name}'")

        # 1. Navigare la Postare
        try:
            if "mbasic" in post_link or "m." in post_link:
                full_link = post_link.replace("mbasic.facebook.com", "www.facebook.com").replace("m.facebook.com", "www.facebook.com")
            else:
                full_link = post_link
            
            self.driver.get(full_link)
            time.sleep(random.uniform(3, 5))
        except Exception as e:
            logging.error(f"❌ Timeout navigare postare: {e}")
            return "TIMEOUT"

        # 2. Click Buton 'Distribuie'
        try:
            share_xpaths = [
                "//div[@aria-label='Send this to friends or post it on your profile.']",
                "//span[contains(text(), 'Distribuie')]",
                "//div[@role='button']//div[contains(text(), 'Distribuie')]",
                "//div[@aria-label='Distribuie']"
            ]
            
            share_btn = None
            for xpath in share_xpaths:
                elems = self.driver.find_elements(By.XPATH, xpath)
                for el in elems:
                    if el.is_displayed():
                        share_btn = el
                        break
                if share_btn: break
            
            if share_btn:
                self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", share_btn)
                time.sleep(1)
                share_btn.click()
                logging.info("ℹ️ Click pe 'Distribuie'.")
                time.sleep(2) 
            else:
                logging.error("❌ Butonul 'Distribuie' nu a fost găsit.")
                return "SHARE_BTN_NOT_FOUND"

        except Exception as e:
            logging.error(f"❌ Eroare la click Distribuie: {e}")
            return "SHARE_BTN_ERROR"

        # 3. Navigare Meniu (Tastatură)
        try:
            found_final_option = False
            logging.info("⌨️ Încep scanarea meniului prin tastatură...")
            
            for i in range(15):
                try:
                    self.actions.send_keys(Keys.TAB).perform()
                    time.sleep(0.3)
                    
                    active_el = self.driver.switch_to.active_element
                    txt = active_el.text.lower()
                    aria = active_el.get_attribute("aria-label")
                    if aria: txt += " " + aria.lower()
                    
                    if "mai multe" in txt or "more options" in txt:
                        logging.info("👉 ENTER pe 'Mai multe opțiuni' (Keyboard).")
                        active_el.send_keys(Keys.ENTER)
                        time.sleep(2) 
                        continue 

                    if "grup" in txt or "group" in txt:
                        logging.info("👉 ENTER pe 'Distribuie într-un grup' (Keyboard).")
                        active_el.send_keys(Keys.ENTER)
                        found_final_option = True
                        break
                except: continue

            if not found_final_option:
                logging.error("❌ Nu am ajuns la opțiunea de grup.")
                return "SHARE_OPTION_FAIL"

            time.sleep(3) 

        except Exception as e:
            logging.error(f"❌ Eroare navigare meniu: {e}")
            return "MENU_NAV_ERROR"

        # 4. CĂUTARE ȘI SELECTARE
        group_selected_successfully = False
        try:
            search_input = self.wait.until(EC.presence_of_element_located((By.XPATH, "//input[@aria-label='Caută grupuri' or @aria-label='Search for groups' or @type='text']")))
            
            import re
            clean_match = re.search(r'^(.*?)(?=\s\d+[\.,]|\s\d{3,})', group_name)
            if clean_match:
                clean_name = clean_match.group(1).strip()
            else:
                clean_name = group_name[:35].strip()

            logging.info(f"🔎 Caut (Nume curat): '{clean_name}'")
            
            search_input.click()
            time.sleep(0.5)
            search_input.clear()
            for char in clean_name:
                search_input.send_keys(char)
                time.sleep(0.05)
            
            logging.info("⏳ Aștept rezultatele (5s)...")
            time.sleep(5) 

            logging.info("🖱️ Caut textul vizibil în listă...")
            
            safe_text = clean_name[:10].replace("'", "").replace('"', '')

            target_xpath = f"//div[@role='dialog']//span[contains(text(), '{safe_text}')]"
            results = self.driver.find_elements(By.XPATH, target_xpath)
            
            if not results:
                 logging.info("⚠️ Nu am găsit textul direct. Caut elemente generice în listă.")
                 results = self.driver.find_elements(By.XPATH, "//div[@role='listbox']//div[@role='option']")

            if not results:
                 results = self.driver.find_elements(By.XPATH, "//div[@role='dialog']//div[contains(@class, 'x1n2onr6')]//div[@role='button']")

            if results:
                target = None
                for res in results:
                    if res.is_displayed():
                        target = res
                        break
                
                if target:
                    try:
                        self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", target)
                        time.sleep(0.5)
                        self.driver.execute_script("arguments[0].click();", target)
                    except:
                        target.click()
                        
                    logging.info("✅ Click realizat pe rezultatul vizual.")
                    group_selected_successfully = True
                else:
                      logging.error("❌ Rezultatele există în DOM dar nu sunt vizibile.")
            else:
                logging.error("❌ Lista de rezultate pare goală (niciun selector nu a mers).")
                self.driver.save_screenshot("debug_empty_list_v2.png")
                self.actions.send_keys(Keys.ESCAPE).perform()
                return "GROUP_NOT_FOUND"

            time.sleep(3)

        except Exception as e:
            logging.error(f"❌ Eroare la căutarea grupului: {e}")
            return "SEARCH_ERROR"

        # 5. POSTARE FINALĂ
        if not group_selected_successfully:
            logging.error("⛔ STOP: Nu postez pentru că nu am selectat niciun grup.")
            return "GROUP_SELECT_FAIL"

        try:
            post_btn_xpath = "//div[@aria-label='Postează' or @aria-label='Post']"
            post_btns = self.driver.find_elements(By.XPATH, post_btn_xpath)
            
            valid_btn = None
            if post_btns:
                for btn in post_btns:
                    if btn.is_displayed():
                        if btn.get_attribute("aria-disabled") == "true":
                            continue
                        valid_btn = btn
                        break
            
            if not valid_btn:
                btns = self.driver.find_elements(By.XPATH, "//div[@role='button']//span[contains(text(), 'Postează') or contains(text(), 'Post')]")
                if btns: valid_btn = btns[0]

            if valid_btn:
                logging.info("✉️ Click 'Postează'...")
                self.driver.execute_script("arguments[0].click();", valid_btn)
                time.sleep(6)
                
                if len(self.driver.find_elements(By.XPATH, "//div[@role='dialog']")) == 0:
                    return "SUCCESS"
                else:
                    self.actions.send_keys(Keys.ESCAPE).perform()
                    return "SUCCESS_POSSIBLE"
            
            logging.error("❌ Nu găsesc butonul final 'Postează' activ.")
            return "POST_BTN_MISSING"

        except Exception as e:
            logging.error(f"❌ Eroare la submit final: {e}")
            return "FINAL_SUBMIT_ERROR"

    def _inject_stealth_scripts(self):
        platform_override = self.selected_config['platform']
        js_payload = f"""
        Object.defineProperty(navigator, 'webdriver', {{ get: () => undefined }});
        Object.defineProperty(navigator, 'platform', {{ get: () => '{platform_override}' }});
        """
        try:
            self.driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {"source": js_payload})
            logging.info("✓ Stealth JS injectat.")
        except: pass

    def _patch_webrtc_leak(self):
        try:
            self.driver.execute_cdp_cmd("Network.enable", {})
            self.driver.execute_cdp_cmd("Network.setBlockedURLs", {"urls": ["*://*/*.webrtc", "*://*/stun*"]})
        except: pass

    def human_typing(self, element, text: str):
        for i, char in enumerate(text):
            if random.random() < 0.05 and i > 0:
                element.send_keys(random.choice('qwertyuiop'))
                time.sleep(random.uniform(0.1, 0.2))
                element.send_keys(Keys.BACK_SPACE)
            element.send_keys(char)
            time.sleep(random.uniform(0.05, 0.15))

    def random_scroll(self, min_px=300, max_px=800):
        try:
            scroll_amount = random.randint(min_px, max_px)
            self.driver.execute_script(f"window.scrollBy(0, {scroll_amount});")
            time.sleep(random.uniform(1, 3))
        except: pass

    def _handle_cookie_consent(self):
        try:
            xpaths = [
                "//button[contains(text(), 'Allow all cookies')]",
                "//button[contains(text(), 'Accept cookies')]",
                "//button[contains(text(), 'Accept')]",
                "//button[contains(text(), 'Permite')]",
                "//div[@aria-label='Allow all cookies']",
                "//span[contains(text(), 'Allow all cookies')]"
            ]
            for xpath in xpaths:
                elements = self.driver.find_elements(By.XPATH, xpath)
                for el in elements:
                    if el.is_displayed():
                        el.click()
                        time.sleep(2)
                        logging.info("🍪 Cookie consent tratat.")
                        return
        except: pass

    def perform_warmup(self):
        logging.info("🔥 Warmup start...")
        try:
            self.driver.get("https://www.facebook.com/")
            time.sleep(5)
            self.random_scroll()
            self.random_scroll()
            logging.info("✓ Warmup finalizat")
        except: pass

# --- 7. ORCHESTRATOR ---
class BotOrchestrator:
    def __init__(self):
        self.db = DatabaseManager()
        self.drive = DriveDownloader()
        self.parser = LocalParser()
        self.bot = None
        self.daily_post_count = 0

    def run(self):
        logging.info("=" * 70)
        logging.info("🤖 FACEBOOK BOT v12.1 - NO TIME LIMITS")
        logging.info("=" * 70)
        try:
            Config.validate()
        except Exception as e:
            logging.critical(f"❌ Config invalid: {e}")
            return

        self.db.cleanup_old_records(days=30)
        while True:
            try:
                self._daily_cycle()
            except KeyboardInterrupt:
                logging.info("\n🛑 Oprire manuală")
                break
            except Exception as e:
                logging.critical(f"🔥 Eroare fatală: {e}")
                time.sleep(300)

    def _daily_cycle(self):
        logging.info("\n" + "=" * 70)
        logging.info("📥 SINCRONIZARE DATE")
        logging.info("=" * 70)
        try:
            self.drive.sync_files()
        except Exception as e:
            logging.error(f"❌ Sync failed: {e}")
            time.sleep(3600)
            return

        base_link = self.parser.get_todays_post()
        raw_groups = self.parser.get_groups()
        if not base_link:
            logging.warning("⚠️ Fără link pentru astăzi. Aștept 1h...")
            time.sleep(3600)
            return
        if not raw_groups:
            logging.warning("⚠️ Fără grupuri. Aștept 1h...")
            time.sleep(3600)
            return

        working_pool = raw_groups[:Config.MAX_GROUPS_POOL]
        pool_size = len(working_pool)
        daily_limit = Config.DAILY_GROUP_LIMIT
        logging.info(f"🎱 Pool activ de grupuri: {pool_size} (limitat la {Config.MAX_GROUPS_POOL})")

        start_index = self.db.get_group_index()
        if start_index >= pool_size:
            start_index = 0
            logging.info("🔄 Index resetat la 0 (era peste limita pool-ului).")

        end_index = start_index + daily_limit
        if end_index <= pool_size:
            todays_groups = working_pool[start_index:end_index]
            next_start_index = end_index
        else:
            remaining_at_end = pool_size - start_index
            overflow_needed = daily_limit - remaining_at_end
            todays_groups = working_pool[start_index:] + working_pool[:overflow_needed]
            next_start_index = overflow_needed
            logging.info("🔄 Am ajuns la capătul listei! Continuăm circular de la început.")

        self.db.update_group_index(next_start_index)
        logging.info(f"\n📋 PLAN ZILNIC (Start: {start_index} -> Next: {next_start_index}):")
        logging.info(f"   Link: {base_link}")
        logging.info(f"   Grupuri selectate: {len(todays_groups)}")

        logging.info("\n🚀 INIȚIALIZARE BOT...")
        try:
            self.bot = FacebookBot()
            if not self.bot.ensure_logged_in():
                logging.error("❌ Login eșuat")
                self.bot.close()
                time.sleep(600)
                return
            if Config.DO_WARMUP:
                self.bot.perform_warmup()
            self.bot.switch_profile()
        except Exception as e:
            logging.error(f"❌ Eroare init bot: {e}")
            return

        # --- MODIFICARE: Fara deadline orar strict ---
        logging.info("\n" + "=" * 70)
        logging.info("📤 START PROCESARE GRUPURI (24/7 MODE)")
        logging.info("=" * 70)

        self.daily_post_count = 0
        for idx, group in enumerate(todays_groups, 1):
            
            # Verificăm doar dacă am depășit limita zilnică (care este implicită prin lista generată)
            # Nu mai verificăm ora curentă.

            if self.db.already_posted_today(group, base_link):
                logging.info(f"[{idx}/{len(todays_groups)}] ⏭️ SKIP (deja postat): {group}")
                continue
            
            if self.bot.is_session_expired():
                logging.info("\n♻️ Sesiune expirată. Refresh...")
                self.bot.refresh_session()
                self.bot.ensure_logged_in()
                self.bot.switch_profile()
            
            logging.info(f"\n[{idx}/{len(todays_groups)}] 📤 PROCESARE: {group}")
            unique_link = generate_unique_link(base_link)
            status = self._post_with_retry(group, unique_link, base_link)
            
            if status in ["SUCCESS", "SUCCESS_FALLBACK"]:
                self.daily_post_count += 1
                logging.info(f"✅ SUCCES (Total: {self.daily_post_count})")
                delay = random.randint(Config.DELAY_MIN_SEC, Config.DELAY_MAX_SEC)
                logging.info(f"⏳ Pauză: {delay}s")
                time.sleep(delay)
            elif status == "NOT_MEMBER":
                logging.warning(f"⚠️ Nu sunt membru")
            else:
                logging.error(f"❌ EȘUAT: {status}")
                time.sleep(30)

        self._show_daily_stats()
        if self.bot:
            self.bot.close()

        # --- MODIFICARE: Calcul Sleep ---
        # Așteptăm până la începutul zilei următoare pentru a respecta limita "Zilnică"
        now = datetime.now()
        tomorrow = now + timedelta(days=1)
        # Setăm următoarea rulare la ora 00:01 a zilei următoare
        next_run = tomorrow.replace(hour=0, minute=1, second=0, microsecond=0)
        
        sleep_seconds = (next_run - now).total_seconds()
        
        # Siguranță: dacă din greșeală iese negativ (ceea ce nu ar trebui), stăm 1 oră
        if sleep_seconds < 0: 
            sleep_seconds = 3600

        logging.info(f"\n😴 Lista completă. SLEEP până mâine la {next_run.strftime('%Y-%m-%d %H:%M')}")
        time.sleep(sleep_seconds)

    def _post_with_retry(self, group_url: str, unique_link: str, base_link: str) -> str:
        for attempt in range(1, Config.MAX_RETRIES + 1):
            try:
                status = self.bot.post_to_group(group_url, unique_link)
                
                if status in ["SUCCESS", "SUCCESS_POSSIBLE", "SUCCESS_FALLBACK"]:
                    self.db.log_post(unique_link, base_link, group_url, "SUCCESS", attempt)
                    return "SUCCESS"
                
                fatal_errors = ["NOT_MEMBER", "GROUP_NOT_FOUND", "SHARE_BTN_NOT_FOUND", "SHARE_OPTION_FAIL"]
                
                if status in fatal_errors:
                    logging.warning(f"⛔ Eroare irecuperabilă ({status}). Nu mai încerc din nou.")
                    self.db.log_post(unique_link, base_link, group_url, status, attempt)
                    return status

                if attempt < Config.MAX_RETRIES:
                    wait_time = (2 ** attempt) * random.uniform(10, 20)
                    logging.info(f"🔄 Retry {attempt}/{Config.MAX_RETRIES} în {wait_time:.1f}s pentru status: {status}")
                    time.sleep(wait_time)
                    self.bot.refresh_session()
                else:
                    self.db.log_post(unique_link, base_link, group_url, f"FAILED_{status}", attempt)
                    return status

            except Exception as e:
                logging.error(f"⚠️ Excepție la tentativa {attempt}: {e}")
                if attempt < Config.MAX_RETRIES:
                    time.sleep(30)
                else:
                    return "ERROR_EXCEPTION"
        
        return "MAX_RETRIES_EXCEEDED"

    def _show_daily_stats(self):
        stats = self.db.get_today_stats()
        logging.info(f"Total: {sum(stats.values())} | {stats}")

# --- 8. ENTRY POINT ---
def main():
    try:
        orchestrator = BotOrchestrator()
        orchestrator.run()
    except KeyboardInterrupt:
        logging.info("\n👋 Oprire gracios...")
    except Exception as e:
        logging.critical(f"❌ Eroare critică: {e}")
        raise
    finally:
        logging.info("🔚 Bot oprit")

if __name__ == "__main__":
    main()